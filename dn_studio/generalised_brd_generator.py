#!/usr/bin/env python3
"""
generalised_brd_generator.py

Standalone version of the notebook's BRD generator that consumes a
structured BRD JSON and writes a DOCX using python-docx only.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, List

from docx import Document  # type: ignore[import]
from docx.enum.style import WD_STYLE_TYPE  # type: ignore[import]


def v(obj: Any, *keys: str, default: str | list | dict | None = "TBD") -> Any:
    cur = obj
    for k in keys:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(k, {})
    if isinstance(cur, dict):
        val = cur.get("value", None)
    else:
        val = cur
    if val is None or val == "" or val == []:
        return default
    return val


def safe_str(val: Any, default: str = "TBD") -> str:
    if val is None or val == "" or (isinstance(val, list) and len(val) == 0):
        return default
    if isinstance(val, list):
        return "; ".join(str(x) for x in val)
    return str(val)


def parse_brd(data: Dict[str, Any]) -> Dict[str, Any]:
    header = data.get("Header", {})
    ctx: Dict[str, Any] = {
        "project_name": safe_str(v(header, "Project Name")),
        "version": safe_str(v(header, "Version")),
        "date": safe_str(v(header, "Date")),
        "business_unit": safe_str(v(header, "Business Unit")),
    }
    return ctx


def build_doc(ctx: Dict[str, Any], dest: str) -> None:
    doc = Document()
    styles = doc.styles
    if "Heading 1" in styles:
        h1 = styles["Heading 1"]
    else:
        h1 = styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
    if "Heading 2" in styles:
        h2 = styles["Heading 2"]
    else:
        h2 = styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)

    doc.add_heading("Business Requirements Document", level=0)
    p = doc.add_paragraph()
    p.add_run(ctx.get("project_name", "Project")).bold = True

    doc.add_paragraph().add_run("Version: ").bold = True
    doc.paragraphs[-1].add_run(ctx.get("version", "1.0"))

    doc.add_paragraph().add_run("Date: ").bold = True
    doc.paragraphs[-1].add_run(ctx.get("date", "TBD"))

    doc.save(dest)
    print(f"  BRD DOCX saved → {dest}")


def main(argv: List[str] | None = None) -> None:
    parser = argparse.ArgumentParser(
        description="Generate BRD DOCX from structured BRD JSON."
    )
    parser.add_argument("input", help="Path to BRD JSON file")
    parser.add_argument(
        "--output", "-o", default="outputs/gcs/brd/BRD.docx", help="Output DOCX path"
    )
    args = parser.parse_args(argv)

    data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    ctx = parse_brd(data)
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    build_doc(ctx, str(out_path))


if __name__ == "__main__":
    main()


